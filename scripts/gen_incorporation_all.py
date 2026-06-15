#!/usr/bin/env python3
"""주식회사 윤희에프엔비 법인 설립 서류 7종 생성 (감사 오병용 반영)"""
import os
import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE_DIR, 'exports', 'contracts')
STATIC_DIR = os.path.join(BASE_DIR, 'app', 'static', 'docs', 'contracts')
os.makedirs(OUT, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

# === 공통 정보 ===
CORP = '주식회사 윤희에프엔비'
REP = '오윤희'
SSN = '781219-2531611'
ADDR = '전라북도 남원시 운봉읍 산덕옛길 24-5'
HQ = '서울특별시 중랑구 봉화산로27길 8, 1층 (중화동)'
HQ_SHORT = '서울특별시 중랑구'
CAPITAL = 10000000
SHARE_PRICE = 5000
TOTAL_SHARES = 10000
ISSUE_SHARES = 2000
PURPOSES = [
    '1. 음식점업',
    '2. 한식 전문점 운영',
    '3. 식품 제조 및 판매업',
    '4. 식품 유통업',
    '5. 프랜차이즈업',
    '6. 식자재 도소매업',
    '7. 부동산 임대업',
    '8. 위 각호에 부대하는 일체의 사업',
]

# === 감사 정보 ===
AUDITOR = '오병용'
AUDITOR_SSN = '760511-1531614'
AUDITOR_ADDR = '서울특별시 노원구 덕릉로 753, 101동 805호 (상계동, 노원롯데캐슬시그니쳐)'

ASCII_ALIASES = {
    f'{CORP}_정관.docx': 'yhfnb_articles.docx',
    f'{CORP}_발기인총회의사록.docx': 'yhfnb_incorporation_minutes.docx',
    f'{CORP}_주식인수증.docx': 'yhfnb_share_subscription.docx',
    f'{CORP}_조사보고서.docx': 'yhfnb_investigation_report.docx',
    f'{CORP}_취임승낙서.docx': 'yhfnb_appointment_acceptance.docx',
    f'{CORP}_주주명부.docx': 'yhfnb_shareholder_registry.docx',
    f'{CORP}_인감신고서.docx': 'yhfnb_seal_registration.docx',
}


def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    return doc


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    p.paragraph_format.space_after = Pt(20)
    return p


def add_p(doc, text, bold=False, size=10, align=None, after=6, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run.bold = bold
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def add_h(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    for i, row_data in enumerate(rows_data):
        for j, text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            if i == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif j == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# 1. 정관
# ============================================================
def gen_articles():
    doc = new_doc()
    add_title(doc, '정    관')

    articles = [
        ('제1장  총  칙', [
            ('제1조 (상호)', f'본 회사는 {CORP}라 칭한다.'),
            ('제2조 (목적)', '본 회사는 다음의 사업을 영위함을 목적으로 한다.\n' + '\n'.join(PURPOSES)),
            ('제3조 (본점소재지)', f'본 회사의 본점은 {HQ_SHORT}에 둔다.'),
            ('제4조 (공고방법)', '본 회사의 공고는 회사의 인터넷 홈페이지에 게재한다. 다만, 전산장애 등의 사유로 게재가 불가능할 때에는 서울특별시에서 발행하는 일간 신문에 게재한다.'),
        ]),
        ('제2장  주  식', [
            ('제5조 (발행할 주식의 총수)', f'본 회사가 발행할 주식의 총수는 {TOTAL_SHARES:,}주로 한다.'),
            ('제6조 (1주의 금액)', f'본 회사가 발행하는 주식 1주의 금액은 금 {SHARE_PRICE:,}원으로 한다.'),
            ('제7조 (설립시에 발행하는 주식의 총수)', f'본 회사가 설립시에 발행하는 주식의 총수는 {ISSUE_SHARES:,}주로 한다.'),
            ('제8조 (주식의 종류)', '본 회사가 발행할 주식은 기명식 보통주식으로 한다.'),
            ('제9조 (주권의 종류)', '본 회사의 주권은 1주권, 5주권, 10주권, 50주권, 100주권, 500주권, 1000주권의 7종으로 한다.'),
            ('제10조 (주식의 양도제한)', '본 회사의 주식을 양도하고자 할 때에는 이사회의 승인을 받아야 한다.'),
            ('제11조 (명의개서대리인)', '① 본 회사는 주식의 명의개서대리인을 둘 수 있다.\n② 명의개서대리인 및 그 사무취급장소와 대행업무의 범위는 이사회의 결의로 정한다.'),
        ]),
        ('제3장  사  채', [
            ('제12조 (사채의 발행)', '본 회사는 이사회의 결의에 의하여 사채를 발행할 수 있다.'),
        ]),
        ('제4장  주주총회', [
            ('제13조 (소집시기)', '① 본 회사의 정기주주총회는 매 사업연도 종료 후 3월 이내에 소집한다.\n② 임시주주총회는 필요에 따라 수시로 이사회의 결의에 의하여 소집한다.'),
            ('제14조 (소집권자)', '주주총회는 법령에 다른 규정이 있는 경우를 제외하고는 대표이사가 소집한다.'),
            ('제15조 (소집통지 및 공고)', '주주총회를 소집함에는 그 일시, 장소 및 회의의 목적사항을 총회일 2주 전에 각 주주에게 서면 또는 전자문서로 통지를 발송하여야 한다.'),
            ('제16조 (의장)', '주주총회의 의장은 대표이사로 한다.'),
            ('제17조 (의결권)', '주주의 의결권은 1주마다 1개로 한다.'),
            ('제18조 (의결권의 대리행사)', '① 주주는 대리인으로 하여금 그 의결권을 행사하게 할 수 있다.\n② 대리인은 주주총회 개시 전에 그 대리권을 증명하는 서면을 제출하여야 한다.'),
            ('제19조 (결의방법)', '주주총회의 결의는 법령에 다른 규정이 있는 경우를 제외하고는 출석한 주주의 의결권의 과반수로 하되 발행주식총수의 4분의 1 이상의 수로 하여야 한다.'),
            ('제20조 (의사록)', '주주총회의 의사에 관하여는 의사록을 작성하고 의장과 출석한 이사가 기명날인 또는 서명하여야 한다.'),
        ]),
        ('제5장  이사와 이사회', [
            ('제21조 (이사의 수)', '본 회사의 이사는 1인 이상 3인 이내로 한다.'),
            ('제22조 (이사의 선임)', '이사는 주주총회에서 선임한다.'),
            ('제23조 (이사의 임기)', '이사의 임기는 3년으로 한다. 다만, 그 임기가 최종의 결산기에 관한 정기주주총회 전에 만료될 경우에는 그 총회의 종결 시까지 그 임기가 연장된다.'),
            ('제24조 (대표이사의 선임)', '대표이사는 주주총회에서 선임한다.'),
            ('제25조 (이사회의 소집)', '① 이사회는 대표이사가 소집한다.\n② 이사회의 소집은 회일 1일 전에 각 이사에게 통지하여야 한다. 다만, 이사 전원의 동의가 있을 때에는 소집절차를 생략할 수 있다.'),
            ('제26조 (이사회의 결의방법)', '이사회의 결의는 이사 과반수의 출석과 출석이사의 과반수로 한다.'),
        ]),
        ('제6장  감  사', [
            ('제27조 (감사의 선임)', f'① 본 회사는 감사 1인을 둔다.\n② 감사는 주주총회에서 선임한다.\n③ 감사의 임기는 취임 후 3년 내의 최종의 결산기에 관한 정기주주총회 종결 시까지로 한다.'),
        ]),
        ('제7장  계  산', [
            ('제28조 (사업연도)', '본 회사의 사업연도는 매년 1월 1일부터 12월 31일까지로 한다.'),
            ('제29조 (재무제표의 작성 등)', '① 대표이사는 매 사업연도 종료 후 다음 각호의 서류와 그 부속명세서 및 영업보고서를 작성하여 이사회의 승인을 받아야 한다.\n  1. 대차대조표\n  2. 손익계산서\n  3. 이익잉여금처분계산서 또는 결손금처리계산서'),
            ('제30조 (이익배당)', '① 이익배당은 금전으로 한다.\n② 이익의 배당을 주주총회의 결의에 의하여 주주에게 지급한다.\n③ 이익배당은 매 결산기말 현재의 주주명부에 기재된 주주 또는 등록된 질권자에게 지급한다.'),
        ]),
        ('제8장  부  칙', [
            ('제31조 (설립비용)', '본 회사의 설립에 소요되는 비용은 금 삼백만원 이내로 한다.'),
            ('제32조 (발기인의 성명과 주소 및 인수주식)', f'본 회사 발기인의 성명, 주민등록번호, 주소 및 인수 주식수는 다음과 같다.\n\n  발기인: {REP}\n  주민등록번호: {SSN}\n  주소: {ADDR}\n  인수 주식수: {ISSUE_SHARES:,}주 (금 {CAPITAL:,}원)'),
            ('제33조 (최초 사업연도)', '본 회사의 최초 사업연도는 회사 설립등기일로부터 2026년 12월 31일까지로 한다.'),
        ]),
    ]

    for chapter_title, items in articles:
        add_h(doc, chapter_title, size=13)
        for art_title, art_body in items:
            add_p(doc, art_title, bold=True, size=11, after=4)
            add_p(doc, art_body, size=10, after=8, indent=0.5)

    add_p(doc, '', after=20)
    add_p(doc, f'위와 같이 {CORP}의 정관을 작성하고 발기인이 기명날인한다.', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, f'발기인   {REP}  (인)', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_정관.docx')
    doc.save(path)
    print(f'  정관: {path}')
    return path


# ============================================================
# 2. 발기인총회의사록 (감사 선임 의안 추가)
# ============================================================
def gen_founders_minutes():
    doc = new_doc()
    add_title(doc, '발기인총회 의사록')

    add_table(doc, [
        ['구 분', '내 용'],
        ['회 사 명', CORP],
        ['일    시', '2026년      월      일      시'],
        ['장    소', HQ],
        ['발기인 총수', '1명'],
        ['출석 발기인', f'1명 ({REP}) — 발기인 전원 출석으로 성립'],
    ], col_widths=[4, 12])

    add_p(doc, '', after=10)
    add_p(doc, f'의장 선출: 출석 발기인 전원의 동의로 발기인 {REP}을(를) 의장으로 선출하다.', after=6)
    add_p(doc, '의장이 개회를 선언하고, 다음의 의안을 상정하여 심의하다.', after=12)

    agendas = [
        ('제1호 의안: 정관 승인의 건',
         f'의장이 미리 작성한 {CORP} 정관 원안을 낭독, 설명하고 승인을 구한 바, 만장일치로 원안대로 승인하다.'),
        ('제2호 의안: 이사 선임의 건',
         f'의장이 설립시 이사 선임을 제안하고, 다음과 같이 이사를 선임하기로 만장일치 가결하다.\n\n  이사: {REP} (주민등록번호: {SSN})'),
        ('제3호 의안: 대표이사 선임의 건',
         f'의장이 대표이사 선임을 제안하고, 다음과 같이 만장일치 가결하다.\n\n  대표이사: {REP}'),
        ('제4호 의안: 감사 선임의 건',
         f'의장이 설립시 감사 선임을 제안하고, 다음과 같이 감사를 선임하기로 만장일치 가결하다.\n\n  감사: {AUDITOR} (주민등록번호: {AUDITOR_SSN})\n  주소: {AUDITOR_ADDR}'),
        ('제5호 의안: 본점소재지 결정의 건',
         f'본점을 {HQ}에 두기로 만장일치 가결하다.'),
        ('제6호 의안: 설립비용 및 발기인 보수의 건',
         '설립에 소요되는 비용은 금 삼백만원 이내로 하며, 발기인 보수는 없는 것으로 만장일치 가결하다.'),
    ]

    for title, body in agendas:
        add_h(doc, title, size=12)
        add_p(doc, body, indent=0.5, after=10)

    add_p(doc, '', after=10)
    add_p(doc, '의장이 이상의 의안이 모두 원안대로 가결되었음을 선언하고 폐회하다.', after=6)
    add_p(doc, '위 결의를 명확히 하기 위하여 의사록을 작성하고 의장 및 출석한 발기인이 기명날인한다.', after=20)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, f'{CORP} 발기인총회', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=15)
    add_p(doc, f'의장 발기인   {REP}  (인)', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_발기인총회의사록.docx')
    doc.save(path)
    print(f'  발기인총회의사록: {path}')
    return path


# ============================================================
# 3. 주식인수증
# ============================================================
def gen_stock_subscription():
    doc = new_doc()
    add_title(doc, '주 식 인 수 증')

    add_p(doc, f'{CORP}의 설립에 있어 발기인으로서 아래와 같이 주식을 인수합니다.', size=11, after=15)

    add_table(doc, [
        ['구 분', '내 용'],
        ['회 사 명', CORP],
        ['1주의 금액', f'금 오천원정 (₩{SHARE_PRICE:,})'],
        ['인수 주식수', f'{ISSUE_SHARES:,}주'],
        ['인수 금액', f'금 일천만원정 (₩{CAPITAL:,})'],
        ['납입 방법', '현금 일시불'],
    ], col_widths=[4, 12])

    add_p(doc, '', after=15)
    add_p(doc, '상기와 같이 주식을 인수하고, 인수 금액 전액을 납입기일까지 납입할 것을 확약합니다.', size=11, after=30)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

    add_p(doc, '주식인수인', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    for line in [
        f'성    명:  {REP}  (인)',
        f'주민등록번호:  {SSN}',
        f'주    소:  {ADDR}',
    ]:
        add_p(doc, line, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

    add_p(doc, '', after=20)
    add_p(doc, f'{CORP} 귀중', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_주식인수증.docx')
    doc.save(path)
    print(f'  주식인수증: {path}')
    return path


# ============================================================
# 4. 조사보고서 (감사 오병용이 작성)
# ============================================================
def gen_investigation_report():
    doc = new_doc()
    add_title(doc, '조 사 보 고 서')

    add_p(doc, f'{CORP}의 설립에 관하여 상법 제298조 및 제313조의 규정에 의거, 감사로서 다음 사항을 조사하고 그 결과를 보고합니다.', size=11, after=15)

    sections = [
        ('1. 정관에 기재된 현물출자 및 재산인수에 관한 사항',
         '  현물출자 사항: 해당 없음\n  재산인수 사항: 해당 없음'),
        ('2. 발기인이 회사설립에 관하여 받을 보수 및 특별이익에 관한 사항',
         '  발기인 보수: 해당 없음\n  발기인 특별이익: 해당 없음'),
        ('3. 회사 부담의 설립비용에 관한 사항',
         '  설립비용: 금 삼백만원 이내 (정관 기재사항과 일치)'),
        ('4. 주식 발행 및 납입에 관한 사항',
         f'  발행할 주식의 총수: {TOTAL_SHARES:,}주\n  1주의 금액: 금 {SHARE_PRICE:,}원\n  설립시 발행주식수: {ISSUE_SHARES:,}주\n  주금 납입액: 금 {CAPITAL:,}원\n  주금 납입은 전액 완료되었음을 확인함'),
    ]

    for title, body in sections:
        add_p(doc, title, bold=True, size=11, after=6)
        add_p(doc, body, size=10, after=12, indent=0.5)

    add_h(doc, '5. 결론', size=12)
    add_p(doc, '위 조사 결과, 상법 제310조에서 정하는 변태설립사항은 존재하지 아니하며, 주식의 인수 및 납입이 적법하게 완료되었음을 확인하고 이에 보고합니다.', size=11, indent=0.5, after=30)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, CORP, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, f'감사   {AUDITOR}  (인)', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_조사보고서.docx')
    doc.save(path)
    print(f'  조사보고서: {path}')
    return path


# ============================================================
# 5. 취임승낙서 (이사 + 대표이사 + 감사)
# ============================================================
def gen_acceptance():
    doc = new_doc()

    # --- 이사 취임승낙서 ---
    add_title(doc, '취 임 승 낙 서')
    add_p(doc, '(이사)', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    add_p(doc, f'본인은 2026년    월    일 개최된 {CORP} 발기인총회에서 이사로 선임되었기에 이를 승낙합니다.', size=11, after=30)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, '이사 취임자', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, f'성    명:  {REP}  (인)', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f'주민등록번호:  {SSN}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f'주    소:  {ADDR}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_p(doc, f'{CORP} 귀중', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    doc.add_page_break()

    # --- 대표이사 취임승낙서 ---
    add_title(doc, '취 임 승 낙 서')
    add_p(doc, '(대표이사)', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    add_p(doc, f'본인은 2026년    월    일 개최된 {CORP} 발기인총회에서 대표이사로 선임되었기에 이를 승낙합니다.', size=11, after=30)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, '대표이사 취임자', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, f'성    명:  {REP}  (인)', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f'주민등록번호:  {SSN}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f'주    소:  {ADDR}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_p(doc, f'{CORP} 귀중', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    doc.add_page_break()

    # --- 감사 취임승낙서 ---
    add_title(doc, '취 임 승 낙 서')
    add_p(doc, '(감사)', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    add_p(doc, f'본인은 2026년    월    일 개최된 {CORP} 발기인총회에서 감사로 선임되었기에 이를 승낙합니다.', size=11, after=30)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, '감사 취임자', size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, f'성    명:  {AUDITOR}  (인)', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f'주민등록번호:  {AUDITOR_SSN}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f'주    소:  {AUDITOR_ADDR}', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_p(doc, f'{CORP} 귀중', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_취임승낙서.docx')
    doc.save(path)
    print(f'  취임승낙서: {path}')
    return path


# ============================================================
# 6. 주주명부
# ============================================================
def gen_shareholder_list():
    doc = new_doc()
    add_title(doc, '주 주 명 부')

    add_p(doc, f'회사명: {CORP}', size=12, bold=True, after=15)

    add_table(doc, [
        ['순번', '주주명', '주민등록번호', '주소', '주식 종류', '주식수', '금액'],
        ['1', REP, SSN, ADDR, '기명식 보통주', f'{ISSUE_SHARES:,}주', f'₩{CAPITAL:,}'],
        ['합계', '', '', '', '', f'{ISSUE_SHARES:,}주', f'₩{CAPITAL:,}'],
    ], col_widths=[1.5, 2.5, 3.5, 4.5, 2.5, 2.0, 2.5])

    add_p(doc, '', after=15)
    add_p(doc, f'발행주식 총수: {ISSUE_SHARES:,}주', size=11, after=6)
    add_p(doc, f'1주의 금액: 금 {SHARE_PRICE:,}원', size=11, after=6)
    add_p(doc, f'자본금 총액: 금 {CAPITAL:,}원', size=11, after=20)

    add_p(doc, '위와 같이 주주명부를 작성합니다.', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, CORP, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, f'대표이사   {REP}  (인)', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_주주명부.docx')
    doc.save(path)
    print(f'  주주명부: {path}')
    return path


# ============================================================
# 7. 인감신고서
# ============================================================
def gen_seal_report():
    doc = new_doc()
    add_title(doc, '인 감 신 고 서')

    add_table(doc, [
        ['구 분', '내 용'],
        ['상    호', CORP],
        ['본    점', HQ],
        ['대표이사', REP],
        ['주민등록번호', SSN],
        ['주    소', ADDR],
    ], col_widths=[4, 12])

    add_p(doc, '', after=20)
    add_p(doc, '위 법인의 대표이사로서 아래의 인감을 법인 인감으로 신고합니다.', size=11, after=20)

    add_p(doc, '법인 인감', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, '┌─────────────────────┐', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '│                                              │', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '│          (인감 날인란)              │', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '│                                              │', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '└─────────────────────┘', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    add_p(doc, '개인 인감 (대표이사)', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_p(doc, '┌─────────────────────┐', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '│                                              │', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '│          (인감 날인란)              │', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '│                                              │', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_p(doc, '└─────────────────────┘', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    add_p(doc, '2026년      월      일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_p(doc, f'신고인   {REP}  (인)', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    path = os.path.join(OUT, f'{CORP}_인감신고서.docx')
    doc.save(path)
    print(f'  인감신고서: {path}')
    return path


# ============================================================
# 실행
# ============================================================
if __name__ == '__main__':
    print('=' * 50)
    print(f'{CORP} 법인 설립 서류 7종 생성 (감사: {AUDITOR})')
    print('=' * 50)
    results = []
    results.append(gen_articles())
    results.append(gen_founders_minutes())
    results.append(gen_stock_subscription())
    results.append(gen_investigation_report())
    results.append(gen_acceptance())
    results.append(gen_shareholder_list())
    results.append(gen_seal_report())
    print('=' * 50)
    print(f'총 {len(results)}종 서류 생성 완료')
    for r in results:
        sz = os.path.getsize(r)
        print(f'  {os.path.basename(r)} ({sz:,} bytes)')
    # 다운로드 경로로 복사
    print('=' * 50)
    print(f'다운로드 경로({STATIC_DIR})로 복사')
    for r in results:
        dst = os.path.join(STATIC_DIR, os.path.basename(r))
        shutil.copy2(r, dst)
        print(f'  -> {dst}')
        alias = ASCII_ALIASES.get(os.path.basename(r))
        if alias:
            alias_dst = os.path.join(STATIC_DIR, alias)
            shutil.copy2(r, alias_dst)
            print(f'  -> {alias_dst}')
    print('=' * 50)
    print('완료')
