#!/usr/bin/env python3
"""인터넷등기소 제출용 법인인감카드 (재)발급신청서 - 별지 제9호 양식 (법무부 예규)"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, shutil

OUT_DIR = "/app/app/static/docs/contracts"
FNAME_KR = "주식회사_윤희에프엔비_법인인감카드신청서.docx"
FNAME_EN = "yhfnb_seal_card_application.docx"
FNAME_ALIAS = "yeoljeong_corp_seal_card_application.docx"

doc = Document()

# --- 페이지 설정 (A4) ---
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            sz, color = val
            el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def set_cell_text(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

# ========== 문서 헤더 ==========
add_para("[별지 제9호 양식]", size=9, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
add_para("인감카드 등 (재)발급신청서", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("(인감의 제출·관리 및 인감증명서 발급에 관한 업무처리지침)", size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# ========== 1. 신청 구분 ==========
t1 = doc.add_table(rows=1, cols=4)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t1)
labels = [("신청 구분", True), ("☑ 신규발급", False), ("☐ 재발급", False), ("☐ 인감증명서\n   발급기능 재부여", False)]
for i, (txt, bold) in enumerate(labels):
    set_cell_text(t1.cell(0, i), txt, size=9, bold=bold)
    if i == 0:
        set_cell_shading(t1.cell(0, i), "E8E8E8")

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ========== 2. 등기기록의 표시 ==========
t2 = doc.add_table(rows=5, cols=2)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2)

# 병합: 첫 행 전체
t2.cell(0, 0).merge(t2.cell(0, 1))
set_cell_text(t2.cell(0, 0), "등 기 기 록 의  표 시", size=10, bold=True)
set_cell_shading(t2.cell(0, 0), "E8E8E8")

rows2 = [
    ("상호(명칭)", "주식회사 윤희에프엔비"),
    ("본점(주사무소)", "서울특별시 중랑구 봉화산로27길 8, 1층 (신내동)"),
    ("등기번호", "(설립등기 완료 후 기재)"),
    ("법인등록번호", "(설립등기 완료 후 기재)"),
]
for i, (label, value) in enumerate(rows2):
    set_cell_text(t2.cell(i+1, 0), label, size=9, bold=True)
    set_cell_shading(t2.cell(i+1, 0), "F5F5F5")
    set_cell_text(t2.cell(i+1, 1), value, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# 컬럼 폭 조정
for row in t2.rows:
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(12.0)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ========== 3. 인감제출자(신청인) ==========
t3 = doc.add_table(rows=7, cols=4)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3)

# 첫 행 병합
t3.cell(0, 0).merge(t3.cell(0, 3))
set_cell_text(t3.cell(0, 0), "인 감 제 출 자 (신 청 인)", size=10, bold=True)
set_cell_shading(t3.cell(0, 0), "E8E8E8")

# 성명/자격
set_cell_text(t3.cell(1, 0), "성  명", size=9, bold=True)
set_cell_shading(t3.cell(1, 0), "F5F5F5")
set_cell_text(t3.cell(1, 1), "오 윤 희", size=9)
set_cell_text(t3.cell(1, 2), "자격(직위)", size=9, bold=True)
set_cell_shading(t3.cell(1, 2), "F5F5F5")
set_cell_text(t3.cell(1, 3), "대표이사", size=9)

# 주민등록번호 (병합)
set_cell_text(t3.cell(2, 0), "주민등록번호", size=9, bold=True)
set_cell_shading(t3.cell(2, 0), "F5F5F5")
t3.cell(2, 1).merge(t3.cell(2, 3))
set_cell_text(t3.cell(2, 1), "781219 - 2******", size=9)

# 주소 (병합)
set_cell_text(t3.cell(3, 0), "주  소", size=9, bold=True)
set_cell_shading(t3.cell(3, 0), "F5F5F5")
t3.cell(3, 1).merge(t3.cell(3, 3))
set_cell_text(t3.cell(3, 1), "전라북도 남원시 운봉읍 산덕옛길 24-5", size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

# 전화번호/이메일
set_cell_text(t3.cell(4, 0), "전화번호", size=9, bold=True)
set_cell_shading(t3.cell(4, 0), "F5F5F5")
set_cell_text(t3.cell(4, 1), "010-3185-6229", size=9)
set_cell_text(t3.cell(4, 2), "이메일", size=9, bold=True)
set_cell_shading(t3.cell(4, 2), "F5F5F5")
set_cell_text(t3.cell(4, 3), "", size=9)

# 인감란 (등기소 제출 인감 + 개인인감)
set_cell_text(t3.cell(5, 0), "등기소에\n제출한 인감", size=9, bold=True)
set_cell_shading(t3.cell(5, 0), "F5F5F5")
set_cell_text(t3.cell(5, 1), "\n\n(인)\n\n", size=9)
set_cell_text(t3.cell(5, 2), "개인인감\n(본인확인)", size=9, bold=True)
set_cell_shading(t3.cell(5, 2), "F5F5F5")
set_cell_text(t3.cell(5, 3), "\n\n(인)\n\n", size=9)

# 인감카드 수령 확인
t3.cell(6, 0).merge(t3.cell(6, 3))
set_cell_text(t3.cell(6, 0), "위와 같이 인감카드 발급을 신청하며, 인감카드를 수령하였음을 확인합니다.", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ========== 4. 비밀번호 ==========
t4 = doc.add_table(rows=3, cols=2)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t4)

t4.cell(0, 0).merge(t4.cell(0, 1))
set_cell_text(t4.cell(0, 0), "비밀번호 (아라비아숫자 6자리)", size=10, bold=True)
set_cell_shading(t4.cell(0, 0), "E8E8E8")

set_cell_text(t4.cell(1, 0), "신규 비밀번호", size=9, bold=True)
set_cell_shading(t4.cell(1, 0), "F5F5F5")
set_cell_text(t4.cell(1, 1), "___    ___    ___    ___    ___    ___", size=11)

set_cell_text(t4.cell(2, 0), "기존 비밀번호\n(재발급 시)", size=9, bold=True)
set_cell_shading(t4.cell(2, 0), "F5F5F5")
set_cell_text(t4.cell(2, 1), "___    ___    ___    ___    ___    ___", size=11)

for row in t4.rows:
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(12.0)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ========== 5. 위임장 ==========
t5 = doc.add_table(rows=5, cols=4)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t5)

t5.cell(0, 0).merge(t5.cell(0, 3))
set_cell_text(t5.cell(0, 0), "위  임  장 (대리인이 신청하는 경우에만 작성)", size=10, bold=True)
set_cell_shading(t5.cell(0, 0), "E8E8E8")

t5.cell(1, 0).merge(t5.cell(1, 3))
set_cell_text(t5.cell(1, 0), "위 본인은 아래 사람을 대리인으로 정하고 인감카드 발급(재발급)에 관한 일체의 권한을 위임합니다.", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

set_cell_text(t5.cell(2, 0), "대리인 성명", size=9, bold=True)
set_cell_shading(t5.cell(2, 0), "F5F5F5")
set_cell_text(t5.cell(2, 1), "", size=9)
set_cell_text(t5.cell(2, 2), "주민등록번호", size=9, bold=True)
set_cell_shading(t5.cell(2, 2), "F5F5F5")
set_cell_text(t5.cell(2, 3), "", size=9)

set_cell_text(t5.cell(3, 0), "주  소", size=9, bold=True)
set_cell_shading(t5.cell(3, 0), "F5F5F5")
t5.cell(3, 1).merge(t5.cell(3, 3))
set_cell_text(t5.cell(3, 1), "", size=9)

set_cell_text(t5.cell(4, 0), "전화번호", size=9, bold=True)
set_cell_shading(t5.cell(4, 0), "F5F5F5")
t5.cell(4, 1).merge(t5.cell(4, 3))
set_cell_text(t5.cell(4, 1), "", size=9)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ========== 유의사항 ==========
add_para("※ 유의사항", size=9, bold=True, space_after=2, space_before=4)
notices = [
    "1. 비밀번호는 생년월일, 주민등록번호, 전화번호 등 추측이 쉬운 번호를 피하십시오.",
    "2. 인감카드는 인감증명서 발급의 핵심 수단이므로 분실하지 않도록 주의하십시오.",
    "3. 인감카드 분실 시 즉시 관할 등기소에 효력정지 신청을 하여야 합니다.",
    "4. 2025년 1월 31일부터 마그네틱 인감카드는 사용이 중단되며, RF카드만 사용 가능합니다.",
    "5. 신규 발급 시 등기소에 제출한 인감을 날인하고, 신분증명서를 지참하여야 합니다.",
    "6. 비밀번호를 3회 이상 잘못 입력하면 카드가 잠기며, 등기소 방문이 필요합니다.",
]
for n in notices:
    add_para(n, size=8, space_after=1)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ========== 날짜 및 서명란 ==========
add_para("2026년       월       일", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para("위 신청인(대리인)                                  (서명 또는 인)", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=16)
add_para("서울중앙지방법원 등기국 귀중", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# ========== 첨부서류 ==========
add_para("", size=6, space_after=2)
add_para("【첨부서류】", size=10, bold=True, space_after=4)
attachments = [
    "1. 신분증명서 사본 (주민등록증, 운전면허증, 여권 중 택1) ··· 1통",
    "2. 개인 인감증명서 (발행 3개월 이내) ··· 1통 (신규 인감 제출 시)",
    "3. 위임장 (대리인 신청 시) ··· 1통",
    "4. 대리인 신분증명서 사본 (대리인 신청 시) ··· 1통",
    "5. 기존 인감카드 (재발급 시 반납)",
]
for att in attachments:
    add_para(att, size=9, space_after=2)

# ========== 저장 ==========
# 한글 파일명
path_kr = os.path.join(OUT_DIR, FNAME_KR)
doc.save(path_kr)

# 영문 별칭 복사
path_en = os.path.join(OUT_DIR, FNAME_EN)
shutil.copy2(path_kr, path_en)

# 기존 별칭도 동기화
path_alias = os.path.join(OUT_DIR, FNAME_ALIAS)
shutil.copy2(path_kr, path_alias)

# 기존 한글 파일명(공백 버전)도 동기화
path_kr_space = os.path.join(OUT_DIR, "주식회사 윤희에프엔비_법인인감카드신청서.docx")
shutil.copy2(path_kr, path_kr_space)

sz = os.path.getsize(path_kr)
print(f"OK: {FNAME_KR} ({sz:,} bytes)")
print(f"OK: {FNAME_EN} (alias)")
print(f"OK: {FNAME_ALIAS} (alias)")
print(f"OK: 공백버전 (alias)")
print(f"Path: {path_kr}")
