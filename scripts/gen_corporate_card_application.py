#!/usr/bin/env python3
"""주식회사 윤희에프엔비 법인카드발급신청서 생성"""
import os
import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE_DIR, 'exports', 'contracts')
STATIC_DIR = os.path.join(BASE_DIR, 'app', 'static', 'docs', 'contracts')
os.makedirs(OUT, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

# === 법인 정보 ===
CORP = '주식회사 윤희에프엔비'
CORP_EN = 'YUNHEE F&B Co., Ltd.'
REP = '오윤희'
SSN = '781219-2531611'
REP_ADDR = '전라북도 남원시 운봉읍 산덕옛길 24-5'
REP_TEL = '010-3185-6229'
HQ = '서울특별시 중랑구 봉화산로27길 8, 1층 (중화동)'
HQ_TEL = '02-               '  # 사업장 전화 (개통 후 기재)
BIZ_REG_NO = '___-__-_____'    # 사업자등록번호 (설립 후 기재)
CORP_REG_NO = '______-_______'  # 법인등록번호 (설립등기 후 기재)
INDUSTRY = '음식점업 (한식 일반음식점)'
ESTABLISHED = '2026년 ___월 ___일'  # 설립등기일
CAPITAL = 10000000
EMPLOYEE_COUNT = '___ 명'

OUT_NAME = f'{CORP}_법인카드발급신청서.docx'
ASCII_ALIAS = 'yhfnb_corporate_card_application.docx'


def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.25
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
    return doc


def set_kor_font(run, name='맑은 고딕'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    set_kor_font(run)
    p.paragraph_format.space_after = Pt(14)
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    set_kor_font(run)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def shade_cell(cell, color='D9E2F3'):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:color="000000"/>'
        '<w:left w:val="single" w:sz="6" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:color="000000"/>'
        '<w:right w:val="single" w:sz="6" w:color="000000"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(tcBorders)


def fill_cell(cell, text, bold=False, size=10, align=None, shade=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    set_kor_font(run)
    set_cell_borders(cell)
    if shade:
        shade_cell(cell, shade)


def build_info_table(doc, rows, label_width=Cm(3.5), value_width=Cm(13.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = label_width
        row.cells[1].width = value_width
        fill_cell(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade='D9E2F3')
        fill_cell(row.cells[1], value)
    return table


def build_doc():
    doc = new_doc()

    # 타이틀
    add_title(doc, '법인카드 발급신청서')

    # 서두 안내
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('당사는 아래와 같이 법인카드 발급을 신청하며, 기재사항이 사실과 다름없음을 확인합니다.')
    run.font.size = Pt(10)
    set_kor_font(run)
    p.paragraph_format.space_after = Pt(10)

    # 1. 법인(신청인) 정보
    add_section_title(doc, '1. 법인(신청인) 정보')
    build_info_table(doc, [
        ('법인명(국문)', CORP),
        ('법인명(영문)', CORP_EN),
        ('법인등록번호', CORP_REG_NO),
        ('사업자등록번호', BIZ_REG_NO),
        ('업    종', INDUSTRY),
        ('설립연월일', ESTABLISHED),
        ('자  본  금', f'금 일천만원정 (₩ {CAPITAL:,})'),
        ('본점 소재지', HQ),
        ('사업장 전화', HQ_TEL),
        ('상시근로자 수', EMPLOYEE_COUNT),
    ])

    # 2. 대표자 정보
    add_section_title(doc, '2. 대표자 정보')
    build_info_table(doc, [
        ('성    명', f'{REP} (인)'),
        ('주민등록번호', SSN),
        ('직    위', '대표이사'),
        ('주    소', REP_ADDR),
        ('휴 대 폰', REP_TEL),
        ('이 메 일', '                                  '),
    ])

    # 3. 신청 카드 정보
    add_section_title(doc, '3. 신청 카드 정보')
    table = doc.add_table(rows=8, cols=4)
    table.autofit = False
    headers = [
        ('카드 종류',  '□ 신용카드   □ 체크카드   □ 기업구매전용카드'),
        ('카드 브랜드', '□ BC   □ 비씨바로   □ 삼성   □ 신한   □ KB국민   □ 하나   □ 우리   □ 농협   □ 기타(        )'),
        ('상품(브랜드명)', '                                                                              '),
        ('월 이용한도(원)', '금                              원정 (₩                    )'),
        ('일 이용한도(원)', '금                              원정 (₩                    )'),
        ('결    제    일', '매월       일 (자동이체)'),
        ('결 제 계 좌',   '은행:                  계좌번호:                              예금주: ' + CORP),
        ('발 급 매 수',   '실물카드        매,   가상카드        매,   해외사용:  □ 허용  □ 차단'),
    ]
    for i, (label, value) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].merge(row.cells[1])
        row.cells[2].merge(row.cells[3])
        # 위 merge로 (a,b) 둘만 남음
        a = row.cells[0]
        b = row.cells[2]
        a.width = Cm(4.5)
        b.width = Cm(12.5)
        fill_cell(a, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade='D9E2F3')
        fill_cell(b, value)

    # 4. 카드 사용(이용) 직원 명단
    add_section_title(doc, '4. 카드 사용(이용) 직원 명단')
    table = doc.add_table(rows=6, cols=5)
    table.autofit = False
    hdr = table.rows[0]
    hdr_titles = ['NO', '성명', '직위', '주민등록번호', '연락처']
    widths = [Cm(1.2), Cm(3.0), Cm(3.0), Cm(4.5), Cm(5.3)]
    for i, t in enumerate(hdr_titles):
        hdr.cells[i].width = widths[i]
        fill_cell(hdr.cells[i], t, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade='D9E2F3')
    # 첫 행: 대표이사
    row1 = table.rows[1]
    data1 = ['1', REP, '대표이사', SSN, REP_TEL]
    for i, t in enumerate(data1):
        row1.cells[i].width = widths[i]
        fill_cell(row1.cells[i], t, align=WD_ALIGN_PARAGRAPH.CENTER)
    # 빈 행
    for r in range(2, 6):
        row = table.rows[r]
        for i in range(5):
            row.cells[i].width = widths[i]
            fill_cell(row.cells[i], str(r) if i == 0 else '', align=WD_ALIGN_PARAGRAPH.CENTER)

    # 5. 제출 서류 체크리스트
    add_section_title(doc, '5. 제출 서류 (체크리스트)')
    docs = [
        '□ 법인카드 발급신청서 (본 서식) 1부',
        '□ 사업자등록증 사본 1부',
        '□ 법인등기부등본(말소사항 포함) 1부 — 발급일 3개월 이내',
        '□ 법인 인감증명서 1부 — 발급일 3개월 이내',
        '□ 사용인감계 1부 (사용인감 사용 시)',
        '□ 대표이사 신분증 사본 1부 (앞·뒷면)',
        '□ 주주명부 1부',
        '□ 위임장 1부 (대리 신청 시) + 대리인 신분증 사본',
        '□ 법인 통장 사본 1부 (결제계좌)',
        '□ 부가가치세 과세표준증명원 또는 재무제표 (한도 심사용, 해당 시)',
    ]
    for d in docs:
        p = doc.add_paragraph()
        run = p.add_run('  ' + d)
        run.font.size = Pt(10)
        set_kor_font(run)
        p.paragraph_format.space_after = Pt(2)

    # 6. 동의 및 확인 사항
    add_section_title(doc, '6. 동의 및 확인 사항')
    consents = [
        '1. 본인(법인)은 위 기재사항이 사실과 다름없음을 확인하며, 허위 기재 시 발생하는 일체의 책임을 부담합니다.',
        '2. 본인(법인)은 신용카드 회원약관 및 개별 상품 약관을 모두 숙지하였으며, 이에 동의합니다.',
        '3. 본인(법인)은 카드 이용대금이 결제일에 지정 계좌에서 자동이체 출금되는 것에 동의합니다.',
        '4. 본인(법인)은 카드 발급 심사 및 한도 관리를 위해 신용정보의 수집·이용·제공·조회에 동의합니다.',
        '5. 카드 사용자(임직원)는 법인 업무 목적으로만 카드를 사용하며, 사적 사용 시 발생한 금액은 사용자가 전액 변제하기로 합니다.',
        '6. 카드 분실·도난·훼손 시 즉시 카드사에 신고하며, 신고 지연으로 발생한 손해는 회원이 부담합니다.',
        '7. 본 신청서에 기재된 개인정보는 카드 발급 및 사후관리 목적으로만 활용되며, 관련 법령에 따라 보관·파기됩니다.',
    ]
    for c in consents:
        p = doc.add_paragraph()
        run = p.add_run(c)
        run.font.size = Pt(9.5)
        set_kor_font(run)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.3)

    # 7. 신청일 및 서명란
    add_section_title(doc, '7. 신청일 및 서명')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('신청일 :  2026년     월     일')
    run.font.size = Pt(11)
    run.bold = True
    set_kor_font(run)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)

    # 서명 테이블
    sig = doc.add_table(rows=3, cols=2)
    sig.autofit = False
    fill_cell(sig.rows[0].cells[0], '신   청   인 (법인)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade='D9E2F3')
    fill_cell(sig.rows[0].cells[1], f'법인명 :  {CORP}', size=11)
    fill_cell(sig.rows[1].cells[0], '대   표   이   사', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade='D9E2F3')
    fill_cell(sig.rows[1].cells[1], f'성  명 :  {REP}                                          (법인인감 印)', size=11)
    fill_cell(sig.rows[2].cells[0], '연   락   처', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade='D9E2F3')
    fill_cell(sig.rows[2].cells[1], f'{REP_TEL}', size=11)

    # 수신처
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run('수신 :   ____________________ 카드(사)   귀중')
    run.font.size = Pt(11)
    run.bold = True
    set_kor_font(run)

    # 안내 (각주)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(
        '※ 본 신청서는 법인 설립등기 완료 및 사업자등록 발급 후 사용 가능합니다.\n'
        '※ 카드사별 양식이 상이할 수 있으며, 실제 신청 시 해당 카드사 공식 양식과 함께 제출하십시오.\n'
        '※ 한도 심사 결과에 따라 신청 금액과 다르게 승인될 수 있습니다.'
    )
    run.font.size = Pt(8.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_kor_font(run)

    return doc


def main():
    doc = build_doc()
    out_path = os.path.join(OUT, OUT_NAME)
    doc.save(out_path)
    print(f'생성 완료 : {out_path}  ({os.path.getsize(out_path):,} bytes)')

    # 정적 경로 복사 (한글명)
    static_path = os.path.join(STATIC_DIR, OUT_NAME)
    shutil.copy2(out_path, static_path)
    print(f'정적 배포   : {static_path}')

    # ASCII alias 복사
    alias_path = os.path.join(STATIC_DIR, ASCII_ALIAS)
    shutil.copy2(out_path, alias_path)
    print(f'ASCII alias : {alias_path}')


if __name__ == '__main__':
    main()
