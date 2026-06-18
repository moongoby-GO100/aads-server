#!/usr/bin/env python3
"""Generate corporate seal card application form (법인인감카드신청서) for court submission."""

import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches


EXPORT_DIR = Path("exports/contracts")
STATIC_DIR = Path("app/static/docs/contracts")

FILENAME = "주식회사 윤희에프엔비_법인인감카드신청서.docx"
ASCII_ALIAS = "yeoljeong_corp_seal_card_application.docx"

DATA = {
    "corp": "주식회사 윤희에프엔비",
    "hq_address": "서울특별시 중랑구 봉화산로27길 8, 1층 (중화동)",
    "corp_reg_no": "(설립등기 후 기재)",
    "rep": "오윤희",
    "rep_rrn": "781219-2531611",
    "rep_address": "전라북도 남원시 운봉읍 산덕옛길 24-5",
    "rep_phone": "010-3185-6229",
    "registry_office": "서울동부지방법원 등기국",
}


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15
    return doc


def add_para(doc, text="", size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(size)
    r.bold = bold
    return p


def set_cell(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(size)
    r.bold = bold
    if fill:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): fill, qn("w:val"): "clear"})
        tc_pr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def merge_cells_in_row(table, row_idx, start_col, end_col):
    cell_start = table.cell(row_idx, start_col)
    cell_end = table.cell(row_idx, end_col)
    cell_start.merge(cell_end)


def set_row_height(row, height_cm):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = tr_pr.makeelement(qn("w:trHeight"), {
        qn("w:val"): str(int(height_cm * 567)),
        qn("w:hRule"): "atLeast"
    })
    tr_pr.append(tr_height)


def build_seal_card_application():
    doc = setup_doc()
    d = DATA

    # Title
    add_para(doc, "", size=6, after=0)
    add_para(doc, "인감카드 발급(갱신) 신청서", size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_para(doc, "(법인용)", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    # Section 1: 발급구분
    add_para(doc, "1. 발급 구분", size=11, bold=True, after=4)

    tbl1 = doc.add_table(rows=1, cols=3)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl1.style = "Table Grid"
    set_cell(tbl1.cell(0, 0), "☑ 신규 발급", size=10, bold=True)
    set_cell(tbl1.cell(0, 1), "☐ 갱 신", size=10)
    set_cell(tbl1.cell(0, 2), "☐ 재발급", size=10)
    set_row_height(tbl1.rows[0], 1.0)

    add_para(doc, "", size=6, after=4)

    # Section 2: 법인 정보
    add_para(doc, "2. 법인 표시", size=11, bold=True, after=4)

    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = "Table Grid"
    set_col_width(tbl2, 0, 4.0)
    set_col_width(tbl2, 1, 13.0)

    labels2 = ["상    호", "본점 소재지", "법인등록번호", "전 화 번 호"]
    values2 = [d["corp"], d["hq_address"], d["corp_reg_no"], d["rep_phone"]]
    for i, (label, value) in enumerate(zip(labels2, values2)):
        set_cell(tbl2.cell(i, 0), label, size=9, bold=True, fill="F2F2F2")
        set_cell(tbl2.cell(i, 1), value, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_row_height(tbl2.rows[i], 0.9)

    add_para(doc, "", size=6, after=4)

    # Section 3: 신청인(대표자) 정보
    add_para(doc, "3. 신청인 (대표자)", size=11, bold=True, after=4)

    tbl3 = doc.add_table(rows=5, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Table Grid"
    set_col_width(tbl3, 0, 4.0)
    set_col_width(tbl3, 1, 13.0)

    labels3 = ["성    명", "주민등록번호", "직    위", "주    소", "전 화 번 호"]
    values3 = [d["rep"], d["rep_rrn"], "대표이사", d["rep_address"], d["rep_phone"]]
    for i, (label, value) in enumerate(zip(labels3, values3)):
        set_cell(tbl3.cell(i, 0), label, size=9, bold=True, fill="F2F2F2")
        set_cell(tbl3.cell(i, 1), value, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_row_height(tbl3.rows[i], 0.9)

    add_para(doc, "", size=6, after=4)

    # Section 4: 인감 날인란
    add_para(doc, "4. 인감 날인", size=11, bold=True, after=4)

    tbl4 = doc.add_table(rows=2, cols=2)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.style = "Table Grid"
    set_col_width(tbl4, 0, 8.5)
    set_col_width(tbl4, 1, 8.5)

    set_cell(tbl4.cell(0, 0), "법인 인감", size=10, bold=True, fill="F2F2F2")
    set_cell(tbl4.cell(0, 1), "대표이사 개인인감", size=10, bold=True, fill="F2F2F2")
    set_row_height(tbl4.rows[0], 0.8)

    set_cell(tbl4.cell(1, 0), "\n\n(인감 날인)\n\n\n", size=9)
    set_cell(tbl4.cell(1, 1), "\n\n(인감 날인)\n\n\n", size=9)
    set_row_height(tbl4.rows[1], 4.5)

    add_para(doc, "", size=6, after=4)

    # Section 5: 비밀번호 설정란
    add_para(doc, "5. 비밀번호 설정", size=11, bold=True, after=4)

    tbl5 = doc.add_table(rows=2, cols=2)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.style = "Table Grid"
    set_col_width(tbl5, 0, 4.0)
    set_col_width(tbl5, 1, 13.0)

    set_cell(tbl5.cell(0, 0), "비밀번호", size=9, bold=True, fill="F2F2F2")
    set_cell(tbl5.cell(0, 1), "(숫자 4자리 — 직접 기재)", size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_row_height(tbl5.rows[0], 0.9)

    set_cell(tbl5.cell(1, 0), "비밀번호 확인", size=9, bold=True, fill="F2F2F2")
    set_cell(tbl5.cell(1, 1), "(동일 번호 재기재)", size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_row_height(tbl5.rows[1], 0.9)

    add_para(doc, "", size=4, after=2)
    add_para(doc, "※ 비밀번호는 인감카드로 인감증명서를 발급받을 때 필요합니다.", size=8,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_para(doc, "※ 비밀번호 분실 시 인감카드를 재발급 받아야 합니다.", size=8,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=6)

    # Section 6: 유의사항
    add_para(doc, "6. 유의사항", size=11, bold=True, after=4)

    notices = [
        "① 인감카드는 법인인감증명서 발급 시 반드시 필요하므로 분실하지 않도록 주의하시기 바랍니다.",
        "② 인감카드를 분실한 경우 즉시 관할 등기소에 분실신고 및 재발급 신청을 하여야 합니다.",
        "③ 인감카드의 비밀번호는 타인에게 노출되지 않도록 관리하여 주십시오.",
        "④ 대리인이 신청하는 경우 위임장과 대리인 신분증을 지참하여야 합니다.",
        "⑤ 신청 시 필요서류: 법인등기부등본, 대표이사 신분증, 법인인감도장, 대표이사 개인인감도장",
    ]
    for notice in notices:
        add_para(doc, notice, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

    add_para(doc, "", size=6, after=8)

    # Section 7: 신청 문구
    add_para(doc,
        "위와 같이 법인 인감카드 발급을 신청합니다.",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    add_para(doc, "2026년       월       일", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    add_para(doc, f"신청인 (대표이사)  {d['rep']}          (인)", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    add_para(doc, "", size=10, after=8)
    add_para(doc, d["registry_office"] + " 귀중", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    return doc


def main():
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    STATIC_DIR.mkdir(parents=True, exist_ok=True)

    doc = build_seal_card_application()
    out_path = EXPORT_DIR / FILENAME
    doc.save(str(out_path))
    print(f"[OK] {out_path}  ({out_path.stat().st_size:,} bytes)")

    static_path = STATIC_DIR / FILENAME
    shutil.copy2(str(out_path), str(static_path))

    alias_path = STATIC_DIR / ASCII_ALIAS
    shutil.copy2(str(out_path), str(alias_path))
    print(f"[OK] static: {static_path}")
    print(f"[OK] alias:  {alias_path}")

    os.system(f'chcon -t httpd_sys_content_t "{static_path}" 2>/dev/null')
    os.system(f'chcon -t httpd_sys_content_t "{alias_path}" 2>/dev/null')
    print("[OK] SELinux context set")


if __name__ == "__main__":
    main()
