from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = BASE_DIR / "exports" / "contracts"
STATIC_DIR = BASE_DIR / "app" / "static" / "docs" / "contracts"
FILENAME = "영업양수도계약서_열정국밥_중화점.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "B7C0CC", size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float = 10.0, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", size: float = 10.0, bold: bool = False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_clause(doc: Document, title: str, items: list[str]) -> None:
    add_paragraph(doc, title, size=11, bold=True)
    for item in items:
        add_paragraph(doc, item, size=9.4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], "1F4E79")
        set_cell_border(header_cells[i])
        for p in header_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=9, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cells[i])
            if i == 0:
                set_cell_shading(cells[i], "EEF3F8")
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=8.5, bold=(i == 0))

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.5)


def build_contract() -> Document:
    doc = Document()
    configure_doc(doc)

    title = add_paragraph(doc, "영 업 양 수 도 계 약 서", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(8)

    add_paragraph(
        doc,
        '양도인 최형섭(이하 "갑")과 양수인 오윤희(이하 "을")는 '
        "열정국밥 중화점의 영업권, 시설, 거래관계 및 영업 운영에 필요한 일체 권리·자산을 "
        "포괄적으로 이전하기 위하여 다음과 같이 계약을 체결한다.",
        size=9.7,
    )

    add_table(
        doc,
        ["구분", "내용"],
        [
            ["대상 영업", "열정국밥 중화점"],
            ["사업자등록번호", "370-37-00939"],
            ["소재지", "서울특별시 중랑구 봉화산로27길 8, 1층 (중화동) / 중화동 90-23 외 1필지 1층 우측점포"],
            ["업태·종목", "음식점업 / 국밥"],
            ["양도 범위", "영업권, 상호 사용 관련 권리, 거래처, 집기·비품 일체, 인테리어, 레시피, 영업 노하우, 임차권 이전 협조, 온라인 채널 운영권한"],
            ["건물 면적", "약 61.28㎡ (철근콘크리트구조)"],
        ],
        widths=[3.2, 14.7],
    )

    add_clause(
        doc,
        "제1조 [계약의 성격 및 포괄양수도]",
        [
            '① 본 계약은 부가가치세법상 사업의 포괄양수도 요건을 충족하는 것을 목적으로 하며, "갑"은 사업장의 인적·물적 조직 및 영업상 필요한 권리·자료를 "을"에게 포괄 이전한다.',
            '② "을"은 개인 명의로 사업자등록을 신청하며, 신청 시 본 계약서를 첨부하여 양수도 사실을 신고한다.',
            '③ "갑"은 "을"의 사업자등록 완료 전까지 임의로 폐업신고를 하지 않으며, "을"의 사업자등록 완료 후 즉시 폐업신고 및 필요한 후속 신고에 협조한다.',
            "④ 세무관청이 포괄양수도에 해당하지 않는다고 판단하여 부가가치세 또는 가산세가 발생하는 경우, 그 귀책사유가 있는 당사자가 이를 부담한다.",
        ],
    )

    add_clause(
        doc,
        "제2조 [양수도 대금 및 지급]",
        [
            "① 권리금은 금 팔천만원정(₩80,000,000)으로 하며, 부가가치세는 별도로 한다.",
            "② 부가가치세가 과세되는 경우 부가가치세는 금 팔백만원정(₩8,000,000)으로 하며, 세금계산서 발행 및 대금 정산은 관련 세법에 따른다.",
            "③ 계약금은 금 팔백만원정(₩8,000,000)으로 하며, 본 계약 체결 전 이미 지급 완료된 것으로 확인한다.",
            "④ 중도금 및 잔금 지급일은 쌍방이 별도 합의하여 아래 지급표에 기재하며, 잔금 지급과 동시에 영업 인도 및 권한 이전을 완료한다.",
        ],
    )

    add_table(
        doc,
        ["구분", "금액", "지급일", "비고"],
        [
            ["계약금", "₩8,000,000", "지급 완료", "기지급 계약금으로 본다"],
            ["중도금", "₩                 ", "2026년     월     일", "쌍방 합의 후 기재"],
            ["잔금", "₩                 ", "2026년     월     일", "영업 인도·서류·권한 이전과 동시 이행"],
            ["VAT", "₩8,000,000", "세금계산서 발행 시", "포괄양수도 인정 시 미부과 가능"],
        ],
        widths=[2.4, 4.1, 4.5, 6.8],
    )

    add_clause(
        doc,
        "제3조 [양도 대상 자산 및 인수인계]",
        [
            '① "갑"은 매장 내 주방기기, 냉장·냉동 설비, 홀 가구, 식기류, 간판, POS 주변기기, CCTV·보안설비, 소모품, 재고, 인테리어 시설물 등 집기·비품 일체를 현 상태 기준으로 "을"에게 양도한다.',
            '② "갑"은 거래처, 식재료·포장재 공급처, 배달대행사, 본사 담당자, 렌탈·리스·보안·통신 업체 연락처 및 계약조건을 인수일까지 목록으로 제공한다.',
            '③ "갑"은 레시피, 조리 프로세스, 영업 노하우, 메뉴·가격표, POS 메뉴 데이터, 배달앱 운영 설정, 광고 설정 및 정산 관련 정보를 성실히 인수인계한다.',
            "④ 양도 대상 집기·비품 및 재고 목록은 별첨으로 작성하고, 인수일에 쌍방 입회하에 최종 확인한다.",
        ],
    )

    add_clause(
        doc,
        "제4조 [인덕션 교체 특약]",
        [
            '① "갑"은 잔금 지급 및 영업 인도 전까지 현재 사용 중인 가스 조리설비를 인덕션 조리기기 8개로 교체 완료하여야 한다.',
            '② 교체 비용, 설치비, 기존 가스설비 철거 또는 안전조치 비용은 전액 "갑"이 부담한다.',
            '③ 교체된 인덕션 8개는 본 계약의 양도 대상 집기·비품에 포함되며, 정상 작동 상태로 인도되어야 한다.',
            '④ 교체 미완료 또는 정상 작동 불가 시 "을"은 잔금 지급 또는 영업 인수를 거절할 수 있고, 이로 인한 지연은 "갑"의 귀책으로 본다.',
        ],
    )

    add_clause(
        doc,
        "제5조 [양도인의 적극 협조 의무]",
        [
            '① "갑"은 "을"의 사업자등록, 영업자 지위승계, 임대차 승계, 본사 가맹계약 승계, 카드·POS·배달앱·온라인 채널 전환에 필요한 서류 제출, 본인확인, 동행, 전자서명 및 연락 응대에 적극 협조한다.',
            '② "갑"은 계약일에 기존 사업자등록증 사본, 최근 6개월 부가세 신고자료, POS 매출자료, 영업신고증 원본 또는 사본, 임대차계약 관련 자료, 행정처분·위생점검 관련 자료, 국세·지방세 완납증명서, 임대료·관리비 완납 확인자료를 제공한다.',
            '③ "갑"은 "을"의 사업자등록 신청 전 본 계약이 체결되어야 함을 확인하고, "을"의 사업자등록 완료 전 폐업신고 금지 의무를 부담하며, 업체정보 삭제, 계정 해지, 정산계좌 임의 변경 등 포괄양수도 요건 또는 영업 연속성을 해치는 행위를 하지 않는다.',
            '④ "갑"은 임대인 최승일로부터 임차권 양도 또는 "을" 명의 신규 임대차계약 체결에 관한 동의를 얻기 위해 "을"과 함께 임대인을 방문하거나 필요한 확인서·동의서 작성에 협조한다.',
            '⑤ "갑"은 인수일 전후 14일 동안 매장 운영 안정화를 위하여 조리, 발주, 배달앱 운영, POS 사용, 정산 확인, 거래처 연결에 관하여 합리적인 범위 내에서 무상 협조한다.',
            '⑥ "갑"은 협조 요청을 받은 때로부터 24시간 이내에 응답하고, 본인 인증·전자서명·방문 동행이 필요한 경우 쌍방이 합의한 일정에 성실히 참석한다.',
            '⑦ "갑"의 협조 지연 또는 거부로 사업자등록, 영업신고, 정산계좌 변경, 온라인 채널 권한 이전이 지연되어 "을"에게 손해가 발생한 경우 "갑"은 그 손해를 배상한다.',
        ],
    )

    add_table(
        doc,
        ["협조 항목", "양도자 적극 협조사항", "완료 시점", "위반 시 효과"],
        [
            ["사업자등록", "사업자등록증 사본, 포괄양수도 확인자료, 사업자등록 신청 첨부자료 제공, 사업자등록 완료 전 폐업신고 금지 및 완료 후 폐업신고", "사업자등록 전후", "VAT·가산세 발생 시 귀책 부담"],
            ["영업신고", "영업신고증 원본 인계, 중랑구청 영업자 지위승계 신고 동행·서명·본인확인", "D-Day", "영업개시 지연 손해 배상"],
            ["임대차", "임대인 동의 동행, 임차권 양도 동의서 또는 을 명의 임대차계약 체결 협조", "계약 전후", "동의 불가 시 해제·계약금 반환"],
            ["본사·가맹", "본사 양도양수 승인 요청, 을 명의 가맹계약 승계, 교육·가맹비 추가 발생 여부 확인", "인수 전", "승계 불가 시 해제 가능"],
            ["온라인 채널", "네이버플레이스 주인 권한 위임, 스마트주문·네이버페이 해지/전환, 카카오맵·구글맵 권한 이전", "사업자등록 직후", "삭제·거부 시 귀책 해제"],
            ["배달·정산", "배달의민족, 쿠팡이츠, 요기요, 땡겨요, 배달대행사 계정·정산계좌 전환, 광고 세팅·리뷰 유지 협조", "D-Day 전후", "정산 누락분 즉시 반환"],
            ["시설·계약", "전기, 수도, 도시가스, 인터넷, 전화, CCTV, 보안, 정수기·렌탈 승계 또는 해지 협조", "D-Day 전후", "숨은 위약금 양도인 부담"],
            ["운영자료", "거래처, 계정·비밀번호, 메뉴·가격, 레시피, 발주처, 포장재, 배달 가능지역, 대행 단가 인수인계", "인수일까지", "미제공 시 잔금 보류 가능"],
            ["직원·보험", "직원 명단, 임금·퇴직금 정산, 4대보험 상태, 화재보험·영업배상보험 승계 가능 여부 고지", "인수 전", "미고지 채무 양도인 부담"],
        ],
        widths=[2.7, 8.2, 3.2, 3.7],
    )

    add_clause(
        doc,
        "제6조 [온라인 플랫폼 및 정산 권한 이전]",
        [
            '① "갑"은 네이버 스마트플레이스 업체 관리 권한을 인수일까지 "을"에게 위임하고, 업체 정보를 임의 삭제·변경하지 않는다.',
            '② "갑"은 네이버 스마트주문, 네이버페이, 카카오맵, 구글맵, SNS, 배달앱 및 기타 온라인 채널의 사업자 정보 변경, 해지, 재가입, 권한 이전 절차에 협조한다.',
            '③ "갑"은 네이버 스마트플레이스의 리뷰·별점 유지를 위해 업체 삭제 후 재등록을 하지 않으며, 주인 권한 위임 방식으로 전환되도록 협조한다.',
            '④ "갑"은 배달앱 광고 설정, 리뷰·평점 승계 가능 여부, 정산계좌 변경 상태를 "을"에게 설명하고, 첫 정산일까지 기존 계정에 입금되는 금액이 있을 경우 즉시 "을"에게 정산한다.',
            '⑤ "갑"은 카드단말기, VAN, POS, 현금영수증, 간편결제, 배달대행 정산계좌가 "을" 명의 통장으로 변경될 수 있도록 필요한 정보를 제공한다.',
            '⑥ 인수일 이후 매출이 "갑" 명의 계정 또는 계좌로 입금된 경우 "갑"은 입금 확인 즉시 "을"에게 통지하고 1영업일 이내에 전액 반환한다.',
        ],
    )

    add_clause(
        doc,
        "제7조 [임대차 및 인허가]",
        [
            '① 현재 임대차 조건은 보증금 금 일천오백만원정(₩15,000,000), 월 차임 금 일백삼십만원정(₩1,300,000, 부가가치세 별도), 관리비 월 금삼만원정(₩30,000)으로 확인한다.',
            '② 임대차보증금은 "을"이 임대인에게 직접 지급하며, 본 계약의 권리금 또는 양수도 대금에 포함되지 않는다.',
            '③ "갑"은 영업자 지위승계 신고, 위생교육, 통신판매업 신고, 음식물폐기물 배출자 신고 등 "을"의 인허가 절차에 필요한 확인 및 서류 제공에 협조한다.',
            '④ 임대인의 서면 동의 또는 "을" 명의 임대차계약 체결이 불가능한 경우 "을"은 본 계약을 해제할 수 있으며, 이 경우 "갑"은 수령한 계약금을 즉시 반환한다.',
        ],
    )

    add_clause(
        doc,
        "제8조 [채무·체납·리스·행정처분 보증]",
        [
            '① "갑"은 인도일 이전에 발생한 국세·지방세 체납, 임대료·관리비 체납, 미지급 거래대금, 직원 임금·퇴직금, 렌탈·리스료, 대출·담보, 과태료 및 행정처분이 없음을 보증한다.',
            '② 인도일 이전 원인으로 발생한 채무, 행정처분, 위생법 위반, 고객 클레임, 세금 및 가산금은 "갑"이 부담한다.',
            '③ "갑"은 국세완납증명서, 지방세완납증명서, 임대료 완납 확인, 거래처 미지급금 확인, 렌탈·리스 잔여계약 현황을 "을"의 요청 시 제출한다.',
            '④ "갑"은 최근 1년 이내 위생점검 결과, 영업정지·과태료·시정명령 등 행정처분 이력, 진행 중인 민원·분쟁·소송·가압류·담보 제공 여부를 계약 체결 전 서면으로 고지한다.',
            '⑤ 숨은 채무 또는 미고지 행정처분으로 "을"에게 손해가 발생한 경우 "갑"은 손해 전액 및 관련 비용을 배상한다.',
        ],
    )

    add_table(
        doc,
        ["확인 항목", "양도인 제출·확인 자료", "계약상 처리"],
        [
            ["세금 체납", "국세완납증명서, 지방세완납증명서", "미제출 또는 체납 발견 시 잔금 보류"],
            ["임대료 체납", "임대인 완납 확인, 관리비 정산 내역", "인도 전 전액 양도인 정산"],
            ["행정처분", "위생점검 결과, 영업정지·과태료 이력", "미고지 시 귀책 해제 및 손해배상"],
            ["미지급 거래대금", "식재료·포장재 거래처 미지급 현황", "인도 전 양도인 부담"],
            ["리스·렌탈", "냉장고, 정수기, 보안, 통신, POS 잔여계약", "승계 합의 없는 잔여금·위약금 양도인 부담"],
        ],
        widths=[3.4, 8.7, 5.7],
    )

    add_clause(
        doc,
        "제9조 [직원·재고·보험]",
        [
            '① 직원 고용승계 여부, 기존 직원의 퇴직금·미지급 임금 정산, 4대보험 상실·취득 신고는 인수일 전 별도 합의서로 정리한다.',
            '② 재고는 인수일에 실사하고, 정상 영업에 사용 가능한 식재료·포장재·소모품만 인수 대상으로 한다.',
            '③ 화재보험, 배상책임보험, CCTV·보안 계약 등은 "을" 명의 신규 가입 또는 승계 여부를 인수일까지 확정한다.',
        ],
    )

    add_clause(
        doc,
        "제10조 [경업금지 및 비밀유지]",
        [
            '① "갑"은 영업 인도일로부터 2년간 서울특별시 중랑구 중화동 및 매장 반경 1km 이내에서 국밥류 등 동종 또는 유사 영업을 직접 또는 제3자를 통하여 영위하지 않는다.',
            '② "갑"은 본 계약 및 인수 과정에서 알게 된 "을"의 개인정보, 정산계좌, 운영자료, 거래처 조건을 제3자에게 누설하지 않는다.',
        ],
    )

    add_clause(
        doc,
        "제11조 [계약 해제 및 위약금]",
        [
            '① "갑"의 귀책사유로 계약이 해제되는 경우 "갑"은 수령한 계약금의 배액인 금 일천육백만원정(₩16,000,000)을 "을"에게 지급한다.',
            '② "을"의 귀책사유로 계약이 해제되는 경우 "을"은 이미 지급한 계약금 금 팔백만원정(₩8,000,000)을 포기한다.',
            '③ 인덕션 8개 교체 미완료, 임대인 동의 협조 거부, 사업자등록·영업신고 협조 거부, 온라인 채널 권한 이전 거부, 숨은 채무 또는 행정처분 미고지는 "갑"의 귀책사유로 본다.',
            '④ 네이버플레이스·배달앱 업체정보 삭제, 리뷰·평점 승계 방해, 사업자등록 전 폐업신고, 기존 정산계좌 입금액 미반환, 본사 가맹승계 또는 임대차 동의 절차 비협조 역시 "갑"의 귀책사유로 본다.',
            '⑤ 잔금 지급기일 이후 일방의 귀책으로 이행이 지체되는 경우 지체일수에 따라 미지급 금액에 연 15%의 비율로 계산한 지연손해금을 지급한다.',
        ],
    )

    add_clause(
        doc,
        "제12조 [특약사항]",
        [
            '1. 본 계약의 양수인은 오윤희 개인이며, "을"은 법인 설립 후 본 계약상 일체의 권리·의무를 설립 법인에 포괄 이전할 수 있고, "갑"은 이에 동의한다.',
            '2. "을"이 법인을 설립한 경우, "을"과 설립 법인은 연대하여 "갑"에 대한 본 계약의 양수도 대금 지급의무를 부담하며, "갑"은 법인 전환에 따른 사업자 정보 변경, 임대차 명의 변경, 온라인 채널·정산 계정 전환 등 필요한 절차에 추가로 협조한다.',
            "3. 개인 명의 계약 후 법인 전환 시 부가가치세법상 포괄양수도 요건(사업의 동일성·계속성) 충족 여부 및 세무 처리는 세무사 검토를 거쳐 처리한다.",
            "4. 권리금은 부가가치세 별도 조건이다. 세금계산서 발행, 포괄양수도 인정 여부 및 VAT 정산은 세무사 검토 후 처리한다.",
            "5. 계약금 금 팔백만원정(₩8,000,000)은 오윤희 개인이 최형섭에게 기지급한 것으로 확인하며, 향후 법인 설립 시 법인 장부에 대표이사 가지급금 또는 법인 자본금으로 정리한다.",
            "6. 별첨 집기·비품 목록, 온라인 채널·계정 인수인계 목록, 거래처·리스·렌탈 현황표는 본 계약과 동일한 효력을 가진다.",
            "7. 본 계약에 명시되지 않은 사항은 민법, 상법, 상가건물 임대차보호법, 부가가치세법 및 일반 상관례에 따른다.",
        ],
    )

    add_paragraph(
        doc,
        "본 계약의 성립을 증명하기 위하여 계약서 2통을 작성하고, 갑과 을이 서명 또는 날인한 후 각 1통씩 보관한다.",
        size=9.7,
    )
    add_paragraph(doc, "2026년        월        일", size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_table(
        doc,
        ["구분", '양도인 "갑"', '양수인 "을"'],
        [
            ["성명", "최형섭", "오윤희"],
            ["주민등록번호", "941209-1082123", "781219-2531611"],
            ["주소", "서울특별시 중랑구 동일로126길 43 (중화동)", "전라북도 남원시 운봉읍 산덕옛길 24-5"],
            ["연락처", "010-9247-5621", "010-3185-6229"],
            ["입금계좌", "[은행명] [계좌번호] 예금주: 최형섭", "-"],
            ["서명/날인", "(인)", "오윤희 (인)"],
        ],
        widths=[3.0, 6.5, 8.3],
    )

    add_paragraph(doc, "별첨 1. 집기·비품 및 재고 실사 목록  /  별첨 2. 거래처·계정·정산채널 인수인계 목록", size=8.7)
    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    STATIC_DIR.mkdir(parents=True, exist_ok=True)

    output = OUTPUT_DIR / FILENAME
    static_output = STATIC_DIR / FILENAME
    if output.exists():
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy2(output, output.with_suffix(output.suffix + f".bak_{stamp}"))

    doc = build_contract()
    doc.save(output)
    shutil.copy2(output, static_output)
    print(output)
    print(static_output)
    print(output.stat().st_size)
    print(static_output.stat().st_size)


if __name__ == "__main__":
    main()
