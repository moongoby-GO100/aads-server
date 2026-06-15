#!/usr/bin/env python3
"""Generate landlord consent and corporation lease drafts for Yeoljeong Gukbap Junghwa."""

import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


EXPORT_DIR = Path("exports/contracts")
STATIC_DIR = Path("app/static/docs/contracts")

LANDLORD_CONSENT = "임대인동의서_주식회사_윤희에프엔비_열정국밥_중화점.docx"
CORP_LEASE = "법인_부동산임대차계약서_열정국밥_중화점.docx"


DATA = {
    "business": "열정국밥 중화점",
    "premises_road": "서울특별시 중랑구 봉화산로27길 8, 1층 (중화동)",
    "premises_lot": "서울특별시 중랑구 중화동 90-23 외 1필지 1층 우측점포",
    "leased_part": "1층 건물 정면 우측점포 약 61.28㎡",
    "use": "음식점(한식, 국밥전문점) 및 법인 본점",
    "deposit": "금 일천오백만원정 (₩15,000,000)",
    "rent": "금 일백삼십만원정 (₩1,300,000, 부가가치세 별도)",
    "maintenance": "월 금삼만원정 (₩30,000)",
    "landlord": "최승일",
    "landlord_rrn": "800722-1031143",
    "landlord_address": "경기도 남양주시 미금로57번길 42, 731동 402호 (다산동)",
    "landlord_phone": "010-4744-2121",
    "landlord_account": "농협 356-1200-005913 (예금주: 최승일)",
    "transferor": "최형섭",
    "transferor_rrn": "941209-1082123",
    "transferor_address": "서울특별시 중랑구 동일로126길 43 (중화동)",
    "transferor_phone": "010-9247-5621",
    "corp": "주식회사 윤희에프엔비",
    "rep": "오윤희",
    "rep_rrn": "781219-2531611",
    "rep_address": "전라북도 남원시 운봉읍 산덕옛길 24-5",
    "rep_phone": "010-3185-6229",
}


def setup_doc(margins=(1.7, 1.5, 1.7, 1.5)):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margins[0])
    section.bottom_margin = Cm(margins[1])
    section.left_margin = Cm(margins[2])
    section.right_margin = Cm(margins[3])

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.1
    return doc


def add_para(doc, text="", size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(size)
    r.bold = bold
    return p


def set_cell(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(size)
    r.bold = bold
    if fill:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): fill, qn("w:val"): "clear"})
        tc_pr.append(shd)


def add_info_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        set_cell(table.cell(i, 0), label, size=8, bold=True, fill="D9EAF7")
        set_cell(table.cell(i, 1), value, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    if widths:
        for row in table.rows:
            row.cells[0].width = Cm(widths[0])
            row.cells[1].width = Cm(widths[1])
    return table


def add_signature_table(doc, parties):
    table = doc.add_table(rows=1 + len(parties), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["구분", "성명/상호", "주민/법인번호", "주소", "연락처", "서명/날인"]
    for j, h in enumerate(headers):
        set_cell(table.cell(0, j), h, size=7.5, bold=True, fill="B4C6E7")
    for i, party in enumerate(parties, start=1):
        for j, value in enumerate(party):
            set_cell(table.cell(i, j), value, size=7.2, align=WD_ALIGN_PARAGRAPH.LEFT if j == 3 else WD_ALIGN_PARAGRAPH.CENTER)
    return table


def save_doc(doc, filename):
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    STATIC_DIR.mkdir(parents=True, exist_ok=True)
    export_path = EXPORT_DIR / filename
    static_path = STATIC_DIR / filename
    doc.save(export_path)
    shutil.copy2(export_path, static_path)
    os.chmod(export_path, 0o644)
    os.chmod(static_path, 0o644)
    return export_path, static_path


def generate_landlord_consent():
    doc = setup_doc(margins=(1.5, 1.3, 1.6, 1.6))
    add_para(doc, "임 대 인  동 의 서", size=17, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(
        doc,
        "본 동의서는 열정국밥 중화점 영업양수도 및 주식회사 윤희에프엔비 설립·사업자등록을 위하여 "
        "임대인이 임차권 승계, 법인 본점 사용 및 법인 명의 신규 임대차계약 체결에 동의하는 문서이다.",
        size=8.5,
        after=6,
    )

    add_para(doc, "1. 부동산 및 임대차 조건", size=10, bold=True, after=3)
    add_info_table(
        doc,
        [
            ("상호/영업장", DATA["business"]),
            ("도로명 주소", DATA["premises_road"]),
            ("지번/임대부분", f'{DATA["premises_lot"]} / {DATA["leased_part"]}'),
            ("용도", DATA["use"]),
            ("보증금", DATA["deposit"]),
            ("월 차임", DATA["rent"]),
            ("관리비", DATA["maintenance"]),
        ],
    )

    add_para(doc, "2. 동의 내용", size=10, bold=True, after=3)
    clauses = [
        f'임대인 {DATA["landlord"]}은 현재 임차인 {DATA["transferor"]}이 위 영업장에 관한 영업권 및 임차권 승계 관련 권리·의무를 {DATA["corp"]}(설립 예정, 대표이사 {DATA["rep"]})에게 이전하는 것에 동의한다.',
        f'임대인은 {DATA["corp"]} 설립등기 및 사업자등록을 위하여 위 부동산을 법인 본점 및 영업장 주소로 사용하는 것에 동의한다.',
        f'임대인은 {DATA["corp"]} 설립등기 완료 후 법인등록번호 및 사업자등록번호를 보완 기재하여 법인 명의 임대차계약을 체결하는 것에 동의한다.',
        f'법인 설립 전 신청·등록 준비행위는 대표자 {DATA["rep"]}이 수행할 수 있으며, 설립 후 임차인 지위는 {DATA["corp"]}에 귀속되도록 한다.',
        "월 차임, 관리비, 공과금, 원상회복, 전대·양도 제한 등 기본 임대차 조건은 별도 법인 임대차계약서에 따른다.",
        "인덕션 8개 교체 및 음식점 영업에 필요한 전기·주방설비 설치·변경은 건물 안전, 관계 법령, 사전 협의 및 필요한 인허가를 전제로 동의한다.",
        "본 동의는 기존 체납 임대료·관리비·공과금이 있는 경우 그 정산 책임을 면제하는 것이 아니며, 발생일 기준 귀책 당사자가 부담한다.",
    ]
    for clause in clauses:
        add_para(doc, f"  - {clause}", size=8.5, after=2)

    add_para(doc, "3. 첨부·확인 권장 서류", size=10, bold=True, after=3)
    add_info_table(
        doc,
        [
            ("필수 확인", "임대차계약서 사본, 임대인 신분증 사본, 건축물대장 또는 등기사항증명서"),
            ("법인 설립 후 보완", "법인등기사항증명서, 법인인감증명서, 사업자등록증"),
            ("영업승계 확인", "영업양수도계약서, 임대료·관리비 완납 확인, 시설 인수인계 확인서"),
        ],
    )

    add_para(doc, "위와 같이 확인하고 동의합니다.", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "2026 년      월      일", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_signature_table(
        doc,
        [
            ("임대인", DATA["landlord"], DATA["landlord_rrn"], DATA["landlord_address"], DATA["landlord_phone"], "(인)"),
            ("현재 임차인", DATA["transferor"], DATA["transferor_rrn"], DATA["transferor_address"], DATA["transferor_phone"], "(인)"),
            ("신규 임차인", f'{DATA["corp"]}\n대표이사 {DATA["rep"]}', "법인등록번호\n[설립 후 기재]", DATA["premises_road"], DATA["rep_phone"], "(인)"),
        ],
    )
    return save_doc(doc, LANDLORD_CONSENT)


def generate_corp_lease():
    doc = setup_doc(margins=(1.2, 1.1, 1.4, 1.4))
    add_para(doc, "법 인  부 동 산  임 대 차  계 약 서", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(
        doc,
        f'※ 본 문서는 {DATA["corp"]} 설립등기 완료 후 법인등록번호·사업자등록번호를 기재하여 체결할 법인 명의 임대차계약서 초안입니다.',
        size=8,
        bold=True,
        after=5,
    )

    add_para(doc, "【부동산의 표시】", size=9.5, bold=True, after=2)
    add_info_table(
        doc,
        [
            ("소재지", DATA["premises_lot"]),
            ("도로명", DATA["premises_road"]),
            ("임대부분", DATA["leased_part"]),
            ("용도", DATA["use"]),
        ],
    )

    add_para(doc, "【계약 내용】", size=9.5, bold=True, after=2)
    add_info_table(
        doc,
        [
            ("보증금", DATA["deposit"]),
            ("월 차임", DATA["rent"]),
            ("관리비", f'{DATA["maintenance"]} / 수도·전기·가스 등 개별공과금은 임차인 부담'),
            ("지급계좌", DATA["landlord_account"]),
            ("기간", "2026년      월      일부터 2028년      월      일까지 (2년)"),
        ],
    )

    add_para(doc, "【일반 약정】", size=9.5, bold=True, after=2)
    clauses = [
        "제1조 [목적] 임대인은 위 부동산을 임차인이 음식점 영업 및 법인 본점으로 사용·수익할 수 있도록 임대하고, 임차인은 보증금과 차임을 지급한다.",
        "제2조 [인도 및 사용] 임대인은 임대차 목적물을 약정한 인도일까지 영업 가능한 상태로 인도하며, 임차인은 선량한 관리자의 주의로 사용한다.",
        "제3조 [차임 지급] 임차인은 월 차임과 부가가치세, 관리비 및 개별공과금을 매월 약정일에 지급한다.",
        "제4조 [전대·양도 제한] 임차인은 임대인의 사전 서면 동의 없이 전대, 임차권 양도 또는 담보 제공을 할 수 없다.",
        "제5조 [수선] 건물 구조, 누수, 기본 설비의 하자는 임대인이 수선하고, 임차인의 귀책 파손 및 영업시설 유지·관리 비용은 임차인이 부담한다.",
        "제6조 [해지] 차임 연체액이 3기의 차임액에 달하거나 중대한 계약 위반이 있는 경우 상대방은 최고 후 계약을 해지할 수 있다.",
        "제7조 [종료 및 원상회복] 계약 종료 시 임차인은 목적물을 원상회복하여 반환하고, 임대인은 연체 차임·손해배상금을 공제한 보증금 잔액을 반환한다.",
        "제8조 [권리금 보호] 임대인은 상가건물 임대차보호법상 임차인의 권리금 회수기회를 정당한 사유 없이 방해하지 않는다.",
        "제9조 [분쟁 해결] 본 계약에 정하지 않은 사항은 민법, 상가건물 임대차보호법 및 일반 상관례에 따른다.",
    ]
    for clause in clauses:
        add_para(doc, clause, size=7.7, after=1)

    add_para(doc, "【특약사항】", size=9.5, bold=True, after=2)
    specials = [
        f'1. 본 계약은 기존 임차인 {DATA["transferor"]}의 영업양수도 및 임대인의 동의에 따라 {DATA["corp"]}가 신규 임차인으로 체결하는 계약이다.',
        f'2. {DATA["corp"]} 설립등기 전에는 대표자 {DATA["rep"]}이 준비행위를 수행하고, 설립등기 완료 후 법인등록번호 및 사업자등록번호를 보완 기재하여 본 계약을 확정한다.',
        "3. 보증금 15,000,000원은 기존 임차인과 신규 임차인의 영업양수도 정산 및 임대인 확인에 따라 처리하며, 임대인이 별도 수령하는 금액이 있는 경우 영수증을 작성한다.",
        "4. 인덕션 8개 교체, 주방·전기 설비 변경, 간판·인테리어 변경은 건물 안전과 관계 법령을 준수하고 임대인과 사전 협의한다.",
        "5. 음식점 영업신고, 사업자등록, 위생교육, 소방·전기 안전 등 영업 관련 인허가와 비용은 임차인이 부담한다.",
        "6. 현 시설 상태로 임대차하며, 임차인은 현장을 확인하였다. 단, 숨은 중대한 하자 또는 구조적 하자는 임대인이 수선한다.",
        "7. 주차 1대 가능, 정화조 비용은 기존 임대차 관행에 따라 임차인 간 분담하여 지급한다.",
        "8. 임차인이 재계약을 원하지 않는 경우 계약만료 3개월 전 임대인에게 사전 고지한다.",
    ]
    for item in specials:
        add_para(doc, item, size=7.7, after=1)

    add_para(doc, "위 계약을 증명하기 위하여 계약 당사자가 서명 또는 기명날인한다.", size=8, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, "2026 년      월      일", size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    add_signature_table(
        doc,
        [
            ("임대인", DATA["landlord"], DATA["landlord_rrn"], DATA["landlord_address"], DATA["landlord_phone"], "(인)"),
            ("임차인", f'{DATA["corp"]}\n대표이사 {DATA["rep"]}', "법인등록번호\n[설립 후 기재]", DATA["premises_road"], DATA["rep_phone"], "(법인인감)"),
        ],
    )
    return save_doc(doc, CORP_LEASE)


def main():
    generated = [
        generate_landlord_consent(),
        generate_corp_lease(),
    ]
    for export_path, static_path in generated:
        print(f"EXPORT {export_path} {export_path.stat().st_size}")
        print(f"STATIC {static_path} {static_path.stat().st_size}")


if __name__ == "__main__":
    main()
