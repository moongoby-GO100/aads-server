#!/usr/bin/env python3
"""
법인인감카드 등 (재)발급신청서 - 별지 제9호 양식
(인감의 제출·관리 및 인감증명서 발급에 관한 업무처리지침, 등기예규 제1661호)
인터넷등기소 제출용
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ('sz', 'val', 'color', 'space'):
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)

def add_paragraph_in_cell(cell, text, bold=False, size=10, align='left'):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'malgun gothic'
    if align == 'center':
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title - 별지 제9호 양식 표기
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('[별지 제9호 양식]')
run.font.size = Pt(9)
run.font.name = 'malgun gothic'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Main title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(6)
run = p.add_run('인감카드 등 (재)발급신청서')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'malgun gothic'

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(18)
run = p.add_run('(인감의 제출·관리 및 인감증명서 발급에 관한 업무처리지침)')
run.font.size = Pt(9)
run.font.name = 'malgun gothic'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ===== Section 1: 신청구분 =====
table1 = doc.add_table(rows=1, cols=4)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = table1.rows[0].cells
add_paragraph_in_cell(cells[0], '신청 구분', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
add_paragraph_in_cell(cells[1], '☑ 신규발급', size=10, align='center')
add_paragraph_in_cell(cells[2], '☐ 재발급', size=10, align='center')
add_paragraph_in_cell(cells[3], '☐ 인감증명서 발급기능 재부여', size=10, align='center')

doc.add_paragraph()

# ===== Section 2: 등기기록 표시 (법인 정보) =====
table2 = doc.add_table(rows=5, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
cells = table2.rows[0].cells
cells[0].merge(cells[3])
add_paragraph_in_cell(cells[0], '등기기록의 표시', bold=True, size=11, align='center')
set_cell_shading(cells[0], 'E8E8E8')

# 상호
cells = table2.rows[1].cells
add_paragraph_in_cell(cells[0], '상호(명칭)', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '주식회사 윤희에프엔비', size=10)

# 본점
cells = table2.rows[2].cells
add_paragraph_in_cell(cells[0], '본점(주사무소)', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '서울특별시 중랑구 봉화산로27길 8, 1층', size=10)

# 등기번호
cells = table2.rows[3].cells
add_paragraph_in_cell(cells[0], '등기번호', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '(설립등기 완료 후 기재)', size=10)

# 법인등록번호
cells = table2.rows[4].cells
add_paragraph_in_cell(cells[0], '법인등록번호', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '(설립등기 완료 후 기재)', size=10)

doc.add_paragraph()

# ===== Section 3: 인감제출자 (신청인) =====
table3 = doc.add_table(rows=6, cols=4)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
cells = table3.rows[0].cells
cells[0].merge(cells[3])
add_paragraph_in_cell(cells[0], '인감제출자(신청인)', bold=True, size=11, align='center')
set_cell_shading(cells[0], 'E8E8E8')

# 성명
cells = table3.rows[1].cells
add_paragraph_in_cell(cells[0], '성  명', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
add_paragraph_in_cell(cells[1], '오윤희', size=10)
add_paragraph_in_cell(cells[2], '자격(직위)', bold=True, size=10, align='center')
set_cell_shading(cells[2], 'F2F2F2')
add_paragraph_in_cell(cells[3], '대표이사', size=10)

# 주민등록번호
cells = table3.rows[2].cells
add_paragraph_in_cell(cells[0], '주민등록번호', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '781219 - 2531611', size=10)

# 주소
cells = table3.rows[3].cells
add_paragraph_in_cell(cells[0], '주  소', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '전라북도 남원시 운봉읍 산덕옛길 24-5', size=10)

# 전화번호
cells = table3.rows[4].cells
add_paragraph_in_cell(cells[0], '전화번호', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
add_paragraph_in_cell(cells[1], '010-3185-6229', size=10)
add_paragraph_in_cell(cells[2], '이메일', bold=True, size=10, align='center')
set_cell_shading(cells[2], 'F2F2F2')
add_paragraph_in_cell(cells[3], '', size=10)

# 인감 날인
cells = table3.rows[5].cells
add_paragraph_in_cell(cells[0], '등기소에\n제출한 인감', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
add_paragraph_in_cell(cells[1], '\n\n        (인)\n\n', size=10, align='center')
add_paragraph_in_cell(cells[2], '개인인감\n(본인확인)', bold=True, size=10, align='center')
set_cell_shading(cells[2], 'F2F2F2')
add_paragraph_in_cell(cells[3], '\n\n        (인)\n\n', size=10, align='center')

doc.add_paragraph()

# ===== Section 4: 비밀번호 =====
table4 = doc.add_table(rows=3, cols=4)
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
cells = table4.rows[0].cells
cells[0].merge(cells[3])
add_paragraph_in_cell(cells[0], '비밀번호 (아라비아숫자 6자리)', bold=True, size=11, align='center')
set_cell_shading(cells[0], 'E8E8E8')

# 신규 비밀번호
cells = table4.rows[1].cells
add_paragraph_in_cell(cells[0], '신규 비밀번호', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '___  ___  ___  ___  ___  ___  (직접 기재)', size=10, align='center')

# 기존 비밀번호 (재발급 시)
cells = table4.rows[2].cells
add_paragraph_in_cell(cells[0], '기존 비밀번호\n(재발급 시)', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '___  ___  ___  ___  ___  ___  (재발급 시에만 기재)', size=10, align='center')

doc.add_paragraph()

# ===== Section 5: 위임장 (대리인 신청 시) =====
table5 = doc.add_table(rows=4, cols=4)
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
cells = table5.rows[0].cells
cells[0].merge(cells[3])
add_paragraph_in_cell(cells[0], '위  임  장 (대리인이 신청하는 경우)', bold=True, size=11, align='center')
set_cell_shading(cells[0], 'E8E8E8')

# 위임 문구
cells = table5.rows[1].cells
cells[0].merge(cells[3])
p = cells[0].paragraphs[0]
p.clear()
run = p.add_run('위 본인은 아래 사람을 대리인으로 정하고 인감카드 발급신청에 관한 일체의 권한을 위임합니다.')
run.font.size = Pt(10)
run.font.name = 'malgun gothic'

# 대리인 정보
cells = table5.rows[2].cells
add_paragraph_in_cell(cells[0], '대리인 성명', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
add_paragraph_in_cell(cells[1], '', size=10)
add_paragraph_in_cell(cells[2], '주민등록번호', bold=True, size=10, align='center')
set_cell_shading(cells[2], 'F2F2F2')
add_paragraph_in_cell(cells[3], '', size=10)

cells = table5.rows[3].cells
add_paragraph_in_cell(cells[0], '주  소', bold=True, size=10, align='center')
set_cell_shading(cells[0], 'F2F2F2')
cells[1].merge(cells[3])
add_paragraph_in_cell(cells[1], '', size=10)

doc.add_paragraph()

# ===== Section 6: 유의사항 =====
p = doc.add_paragraph()
p.space_before = Pt(6)
run = p.add_run('※ 유의사항')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'malgun gothic'

notes = [
    '1. 비밀번호는 생년월일, 주민등록번호, 전화번호 등 추측이 쉬운 번호를 피하십시오.',
    '2. 인감카드는 인감증명서 발급의 핵심 수단이므로 분실하지 않도록 주의하십시오.',
    '3. 인감카드 분실 시 즉시 관할등기소에 효력정지 신청을 하여야 합니다.',
    '4. 2025년 1월 31일부터 마그네틱 인감카드는 사용이 중단되며, RF카드만 사용 가능합니다.',
    '5. 신규 발급 시 등기소에 제출한 인감을 날인하고, 신분증명서를 지참하여야 합니다.',
]
for note in notes:
    p = doc.add_paragraph()
    run = p.add_run(note)
    run.font.size = Pt(8.5)
    run.font.name = 'malgun gothic'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.space_after = Pt(2)

doc.add_paragraph()

# ===== Section 7: 신청일 및 제출처 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(12)
run = p.add_run('2026년       월       일')
run.font.size = Pt(11)
run.font.name = 'malgun gothic'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(24)
run = p.add_run('위 신청인(대리인)                                  (서명 또는 인)')
run.font.size = Pt(11)
run.font.name = 'malgun gothic'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(30)
run = p.add_run('서울중앙지방법원 등기국 귀중')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'malgun gothic'

# ===== Section 8: 첨부서류 안내 =====
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('【첨부서류】')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'malgun gothic'

attachments = [
    '1. 신분증명서 사본 (주민등록증, 운전면허증, 여권 중 택1) ··· 1통',
    '2. 인감증명서 (개인, 발행 3개월 이내) ··· 1통 (신규제출 시)',
    '3. 위임장 (대리인 신청 시) ··· 1통',
    '4. 대리인 신분증명서 사본 (대리인 신청 시) ··· 1통',
]
for att in attachments:
    p = doc.add_paragraph()
    run = p.add_run(att)
    run.font.size = Pt(9)
    run.font.name = 'malgun gothic'
    p.space_after = Pt(2)

# Save
output_dir = '/app/app/static/docs/contracts'
filename = '주식회사 윤희에프엔비_법인인감카드신청서.docx'
filepath = os.path.join(output_dir, filename)
doc.save(filepath)

# Also save ASCII alias
alias = 'yeoljeong_corp_seal_card_application.docx'
alias_path = os.path.join(output_dir, alias)
doc.save(alias_path)

print(f"OK: {filepath} ({os.path.getsize(filepath)} bytes)")
print(f"OK: {alias_path} ({os.path.getsize(alias_path)} bytes)")
